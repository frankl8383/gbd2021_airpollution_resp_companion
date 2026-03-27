#!/usr/bin/env python3
"""Build submission-ready docx files and a packaged submission folder."""

from __future__ import annotations

import csv
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
WRITING = PROJECT_ROOT / "writing"
OUTPUT_MAIN = PROJECT_ROOT / "outputs" / "main"
OUTPUT_SUPP = PROJECT_ROOT / "outputs" / "supp"
LOGS = PROJECT_ROOT / "logs"
DOCS = PROJECT_ROOT / "docs"
SUBMISSION_ROOT = PROJECT_ROOT / "submission_packages"
PUBLIC_COMPANION_URL = "https://github.com/frankl8383/gbd2021_airpollution_resp_companion"
ZENODO_DOI = "10.5281/zenodo.19222088"

PACKAGE_SPECS = [
    {
        "upload_prefix": "BMC_submission_package",
        "repro_prefix": "BMC_reproducibility_bundle",
        "current_upload_name": "BMC_submission_package_current",
        "current_repro_name": "BMC_reproducibility_bundle_current",
        "journal_name": "BMC Public Health",
        "cover_letter_source": WRITING / "cover_letter_bmc_v1.md",
    },
]

TABLE_COLUMN_WIDTHS = {
    "table2_breakpoint_main.csv": [0.55, 0.48, 0.58, 0.72, 0.52, 0.58, 0.78, 0.58, 0.78, 0.46],
    "table3_decomposition_main.csv": [0.60, 0.48, 0.56, 0.72, 0.58, 0.60, 0.56, 0.78, 0.44],
}


def clean_inline_md(text: str) -> str:
    text = text.rstrip()
    text = re.sub(r"\\([*_`])", r"\1", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\^([^^]+)\^", r"\1", text)
    text = re.sub(r"^\*\s+", "", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()


def set_doc_style(doc: Document, *, font_size: float = 12, line_spacing: float = 2.0) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(font_size)
    normal.paragraph_format.line_spacing = line_spacing
    normal.paragraph_format.space_after = Pt(0)


def enable_line_numbering(doc: Document) -> None:
    sect_pr = doc.sections[0]._sectPr
    existing = sect_pr.xpath("./w:lnNumType")
    if existing:
        return
    ln_num_type = OxmlElement("w:lnNumType")
    ln_num_type.set(qn("w:countBy"), "1")
    ln_num_type.set(qn("w:restart"), "newPage")
    sect_pr.append(ln_num_type)


def add_page_number_field(paragraph) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def add_footer_page_number(doc: Document) -> None:
    footer = doc.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Page ")
    add_page_number_field(paragraph)


def set_run_font(run, *, size: float | None = None, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_paragraph(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    size: float | None = None,
    align: str | None = None,
    line_spacing: float = 2.0,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(0)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)


def parse_section_block(text: str, header: str) -> str:
    pattern = re.compile(rf"^## {re.escape(header)}\n\n", re.M)
    match = pattern.search(text)
    if not match:
        return ""
    start = match.end()
    next_match = re.search(r"^## ", text[start:], re.M)
    end = start + next_match.start() if next_match else len(text)
    return text[start:end].strip()


def next_version_name(prefix: str) -> str:
    SUBMISSION_ROOT.mkdir(parents=True, exist_ok=True)
    pattern = re.compile(rf"^{re.escape(prefix)}_v(\d+)$")
    versions = []
    for path in SUBMISSION_ROOT.iterdir():
        if not path.is_dir():
            continue
        match = pattern.match(path.name)
        if match:
            versions.append(int(match.group(1)))
    next_version = max(versions, default=0) + 1
    return f"{prefix}_v{next_version}"


def parse_caption_sections() -> dict[str, list[str]]:
    lines = (WRITING / "figure_table_captions.md").read_text(encoding="utf-8").splitlines()
    sections: dict[str, list[str]] = {}
    current: str | None = None
    for line in lines:
        stripped = line.strip()
        if not stripped or stripped.startswith("# "):
            continue
        if stripped.startswith("## "):
            current = clean_inline_md(stripped[3:])
            sections[current] = []
            continue
        if current is not None:
            sections[current].append(clean_inline_md(stripped))
    return sections


def get_caption_lines(section_name: str, sections: dict[str, list[str]]) -> list[str]:
    return [line for line in sections.get(section_name, []) if line]


def add_caption_block(doc: Document, section_name: str, sections: dict[str, list[str]]) -> None:
    lines = get_caption_lines(section_name, sections)
    if not lines:
        return
    add_paragraph(doc, lines[0], bold=True, size=12)
    for line in lines[1:]:
        add_paragraph(doc, line)


def _append_trpr_flag(row, tag: str) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.xpath(f"./w:{tag}")
    if existing:
        return
    flag = OxmlElement(f"w:{tag}")
    flag.set(qn("w:val"), "true")
    tr_pr.append(flag)


def _set_cell_borders(
    cell,
    *,
    top: str = "nil",
    bottom: str = "nil",
    left: str = "nil",
    right: str = "nil",
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge, value in {
        "top": top,
        "bottom": bottom,
        "left": left,
        "right": right,
    }.items():
        border = tc_borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            tc_borders.append(border)
        border.set(qn("w:val"), value)
        if value != "nil":
            border.set(qn("w:sz"), "8")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")


def add_word_table(doc: Document, csv_path: Path) -> None:
    with csv_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.reader(handle))
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = False
    widths = TABLE_COLUMN_WIDTHS.get(csv_path.name, [])

    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = str(value)
            if col_idx < len(widths):
                cell.width = Inches(widths[col_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 0.92
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run_font(run, size=8.4 if row_idx == 0 else 8.0, bold=(row_idx == 0))
            if row_idx == 0:
                _set_cell_borders(cell, top="single", bottom="single")
            elif row_idx == len(rows) - 1:
                _set_cell_borders(cell, bottom="single")
            else:
                _set_cell_borders(cell)
        _append_trpr_flag(table.rows[row_idx], "cantSplit")
        if row_idx == 0:
            _append_trpr_flag(table.rows[row_idx], "tblHeader")


def parse_title_page_source() -> dict[str, str]:
    text = (WRITING / "manuscript_full_draft_v10.md").read_text(encoding="utf-8")
    return {
        "title": parse_section_block(text, "Title").splitlines()[0].replace("**", "").strip(),
        "running_title": parse_section_block(text, "Running Title").splitlines()[0].replace("**", "").strip(),
        "authors": parse_section_block(text, "Authors").splitlines()[0].replace("**", "").replace("^", "").strip(),
        "affiliations": parse_section_block(text, "Affiliations"),
        "emails": parse_section_block(text, "Author Emails"),
        "correspondence": parse_section_block(text, "Correspondence"),
        "keywords": parse_section_block(text, "Keywords"),
    }


def build_title_page_docx(out_path: Path) -> None:
    meta = parse_title_page_source()
    doc = Document()
    set_doc_style(doc)
    add_paragraph(doc, meta["title"], bold=True, size=14, align="center")
    add_paragraph(doc, meta["authors"], align="center")
    add_paragraph(doc, "", align="center")
    add_paragraph(doc, "Affiliations", bold=True, size=12)
    for line in meta["affiliations"].splitlines():
        if line.strip():
            add_paragraph(doc, clean_inline_md(line))
    add_paragraph(doc, "Author emails", bold=True, size=12)
    for line in meta["emails"].splitlines():
        if line.strip():
            add_paragraph(doc, clean_inline_md(line))
    add_paragraph(doc, "Correspondence", bold=True, size=12)
    add_paragraph(doc, clean_inline_md(meta["correspondence"]))
    add_paragraph(doc, "Running title", bold=True, size=12)
    add_paragraph(doc, meta["running_title"])
    add_paragraph(doc, "Keywords", bold=True, size=12)
    for line in meta["keywords"].splitlines():
        if line.strip().startswith("- "):
            add_paragraph(doc, clean_inline_md(line[2:]))
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def build_cover_letter_docx(source_path: Path, out_path: Path) -> None:
    text = source_path.read_text(encoding="utf-8").strip().splitlines()
    doc = Document()
    set_doc_style(doc, font_size=10.5, line_spacing=1.05)
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    for line in text:
        stripped = line.strip()
        if not stripped:
            add_paragraph(doc, "", size=10.5, line_spacing=1.05)
            continue
        if stripped.startswith("# "):
            continue
        add_paragraph(doc, clean_inline_md(stripped), size=10.5, line_spacing=1.05)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def build_manuscript_docx(out_path: Path, *, repro_bundle_filename: str) -> None:
    text = (WRITING / "manuscript_full_draft_v10.md").read_text(encoding="utf-8")
    text = text.replace("__REPRO_BUNDLE_FILENAME__", repro_bundle_filename)
    text = text.replace("BMC_reproducibility_bundle_current.zip", repro_bundle_filename)
    caption_sections = parse_caption_sections()
    doc = Document()
    set_doc_style(doc)
    enable_line_numbering(doc)
    add_footer_page_number(doc)

    skip_sections = {
        "Running Title",
        "Authors",
        "Affiliations",
        "Author Emails",
        "Correspondence",
    }
    current_section = ""
    skip_section = False

    lines = text.splitlines()
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            continue
        if stripped.startswith("## "):
            current_section = clean_inline_md(stripped[3:])
            skip_section = current_section in skip_sections
            if current_section == "Title":
                skip_section = False
                continue
            if skip_section:
                continue
            add_paragraph(doc, current_section, bold=True, size=13)
            continue
        if skip_section:
            continue
        if current_section == "Title":
            add_paragraph(doc, clean_inline_md(stripped), bold=True, size=14, align="center")
            current_section = ""
            continue
        if stripped.startswith("### "):
            subtitle = clean_inline_md(stripped[4:])
            add_paragraph(doc, subtitle, bold=True, size=12)
            continue
        if stripped.startswith("- "):
            add_paragraph(doc, clean_inline_md(stripped[2:]))
            continue
        add_paragraph(doc, clean_inline_md(stripped))

    add_paragraph(doc, "Figure and table captions", bold=True, size=13)
    for figure_name in ["Figure 1", "Figure 2", "Figure 3", "Figure 4"]:
        add_caption_block(doc, figure_name, caption_sections)
    for table_name, csv_path in [
        ("Table 1", OUTPUT_MAIN / "table2_breakpoint_main.csv"),
        ("Table 2", OUTPUT_MAIN / "table3_decomposition_main.csv"),
    ]:
        lines = get_caption_lines(table_name, caption_sections)
        if not lines:
            continue
        add_paragraph(doc, lines[0], bold=True, size=12)
        add_word_table(doc, csv_path)
        for legend_line in lines[1:]:
            add_paragraph(doc, legend_line)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def copy_file(src: Path, dst_dir: Path) -> None:
    dst_dir.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst_dir / src.name)


def copy_relative_file(src: Path, dst_root: Path) -> None:
    rel = src.relative_to(PROJECT_ROOT)
    dst = dst_root / rel
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(src, dst)


def write_additional_files_manifest(dst: Path, repro_bundle_name: str) -> None:
    text = f"""# Additional Files Upload Map

- Additional file 1: `breakpoint_model_comparison_1990_2021.csv` (CSV). Breakpoint model comparison output for the main `1990–2021` analysis.
- Additional file 2: `breakpoint_model_comparison_1990_2019.csv` (CSV). Breakpoint model comparison output for the `1990–2019` sensitivity analysis.
- Additional file 3: `four_factor_decomposition_all_periods_by_sex.csv` (CSV). Sex-stratified four-factor decomposition output for the main and sensitivity windows, including the full, pre-break, and post-break analysis intervals.
- Additional file 4: `four_factor_decomposition_capped_paf_both.csv` (CSV). Both-sex four-factor decomposition rerun using capped implied PAF values.
- Additional file 5: `four_factor_decomposition_annual_chain_both.csv` (CSV). Both-sex four-factor decomposition rerun using annual-chain accumulation.
- Additional file 6: `four_factor_decomposition_common_age_set_both.csv` (CSV). Full both-sex four-factor decomposition rerun table on the common age set, with substantive differences confined to COPD deaths after excluding structurally empty `<15` cells.
- Additional file 7: `four_factor_decomposition_sensitivity_both.csv` (CSV). Both-sex four-factor decomposition rerun for the `1990–2019` window.
- Additional file 8: `vulnerability_focus_long.csv` (CSV). Long-format vulnerability summary for under-5 and age-70-plus burden shares.
- Additional file 9: `vulnerability_peak_age_long.csv` (CSV). Long-format peak-age vulnerability output.
- Additional file 10: `{repro_bundle_name}` (ZIP). Audit and manuscript-packaging bundle containing the current source files, breakpoint-ready summary input, manuscript-facing output tables and figures, QC logs, a run-status manifest, and scripts required to rerun breakpoint analyses and rebuild the manuscript-facing tables, figures, and submission package. Complete harmonization, four-factor decomposition, and vulnerability reruns require upstream inputs that are not redistributed in the journal bundle.
"""
    (dst / "ADDITIONAL_FILES_UPLOAD_MAP.md").write_text(text, encoding="utf-8")


def write_repro_bundle_run_status(dst: Path) -> None:
    text = """# Run Status

This manifest distinguishes between scripts that can be rerun directly from the bundled inputs and scripts retained mainly for audit of the manuscript workflow.

## Directly Rerunnable From Bundled Inputs

- `scripts/00_validate_gbd_downloads.py`
- `scripts/30_breakpoint_trends.py`
- `scripts/60_make_publication_tables.py`
- `scripts/61_make_publication_figures.py`
- `scripts/70_build_submission_package.py`

## Retained For Audit, But Not Fully Rerunnable From Bundled Inputs Alone

- `scripts/20_harmonize_national_gbd.py`
- `scripts/21_prepare_four_factor_inputs.py`
- `scripts/41_decompose_national_burden_four_factor.py`
- `scripts/50_profile_vulnerability.py`

`scripts/00_validate_gbd_downloads.py` is included as an audit/validation helper; in the distributed bundle it is expected to report missing upstream raw-download files rather than a fully satisfied raw-data environment.

These scripts require upstream raw or intermediate inputs that are not redistributed in the journal bundle. The bundled files are intended to support breakpoint rerun plus manuscript-facing table, figure, and submission-package regeneration.

## Internal Output To Manuscript Mapping

- `outputs/main/table2_breakpoint_main.csv` -> manuscript Table 1
- `outputs/main/table3_decomposition_main.csv` -> manuscript Table 2

## Rebuild Label Note

In a clean unzip environment, regenerated package labels restart from `v1/current`. These regenerated labels are non-archival convenience outputs. The archived peer-review snapshot is the frozen package pair distributed with this release.
"""
    (dst / "RUN_STATUS.md").write_text(text, encoding="utf-8")


def build_upload_bundle(
    bundle_dir: Path,
    package_name: str,
    journal_name: str,
    cover_letter_source: Path,
    repro_bundle_zip: Path,
) -> None:
    if bundle_dir.exists():
        shutil.rmtree(bundle_dir)
    bundle_dir.mkdir(parents=True, exist_ok=True)

    manuscript_dir = bundle_dir / "manuscript_files"
    figures_dir = bundle_dir / "figures"
    supplemental_dir = bundle_dir / "supplementary_analytical_files"

    build_manuscript_docx(
        manuscript_dir / "01_submission_manuscript_v1.docx",
        repro_bundle_filename=repro_bundle_zip.name,
    )
    build_title_page_docx(manuscript_dir / "02_title_page_v1.docx")
    build_cover_letter_docx(cover_letter_source, manuscript_dir / "03_cover_letter_v1.docx")

    figure_files = [
        OUTPUT_MAIN / "figure1_study_workflow.png",
        OUTPUT_MAIN / "figure1_study_workflow.pdf",
        OUTPUT_MAIN / "figure2_trends_breakpoints.png",
        OUTPUT_MAIN / "figure2_trends_breakpoints.pdf",
        OUTPUT_MAIN / "figure3_decomposition.png",
        OUTPUT_MAIN / "figure3_decomposition.pdf",
        OUTPUT_MAIN / "figure4_vulnerability.png",
        OUTPUT_MAIN / "figure4_vulnerability.pdf",
    ]
    for src in figure_files:
        copy_file(src, figures_dir)

    additional_files = [
        OUTPUT_SUPP / "breakpoint_model_comparison_1990_2021.csv",
        OUTPUT_SUPP / "breakpoint_model_comparison_1990_2019.csv",
        OUTPUT_SUPP / "four_factor_decomposition_all_periods_by_sex.csv",
        OUTPUT_SUPP / "four_factor_decomposition_capped_paf_both.csv",
        OUTPUT_SUPP / "four_factor_decomposition_annual_chain_both.csv",
        OUTPUT_SUPP / "four_factor_decomposition_common_age_set_both.csv",
        OUTPUT_SUPP / "four_factor_decomposition_sensitivity_both.csv",
        OUTPUT_SUPP / "vulnerability_focus_long.csv",
        OUTPUT_SUPP / "vulnerability_peak_age_long.csv",
    ]
    for src in additional_files:
        copy_file(src, supplemental_dir)

    copy_file(repro_bundle_zip, supplemental_dir)
    write_additional_files_manifest(supplemental_dir, repro_bundle_zip.name)
    write_upload_package_readme(
        bundle_dir,
        package_name,
        journal_name,
        grouped=True,
        repro_bundle_name=repro_bundle_zip.name,
    )


def build_current_repro_bundle(
    repro_sources: list[Path],
    current_repro_dir: Path,
    bundle_name: str,
    *,
    current: bool,
) -> None:
    if current_repro_dir.exists():
        shutil.rmtree(current_repro_dir)
    current_repro_dir.mkdir(parents=True, exist_ok=True)

    for src in repro_sources:
        copy_relative_file(src, current_repro_dir)

    write_repro_bundle_readme(current_repro_dir, bundle_name, current=current)
    write_repro_bundle_run_status(current_repro_dir)


def write_upload_package_readme(
    dst: Path,
    package_name: str,
    journal_name: str,
    *,
    grouped: bool = False,
    repro_bundle_name: str = "",
) -> None:
    version_label = package_name.rsplit("_", 1)[-1]
    if version_label == "current":
        summary_line = (
            "This package is a non-archival convenience output for the current journal-upload "
            "bundle. Rebuilding the submission package updates this directory in place."
        )
    else:
        summary_line = (
            f"This package contains the journal-upload materials prepared for `{journal_name}`."
        )
    manuscript_prefix = "manuscript_files/" if grouped else ""
    figure_prefix = "figures/" if grouped else ""
    supplemental_prefix = "supplementary_analytical_files/" if grouped else ""
    resolved_repro_bundle_name = repro_bundle_name or (
        package_name.replace("submission_package", "reproducibility_bundle") + ".zip"
    )
    text = f"""# {journal_name} Submission Package {version_label}

{summary_line}

Upload the files below individually in the journal portal. For figures, upload the PDF version only; the PNG copy is kept in this package for local preview.

## Main Files

- Main manuscript: `{manuscript_prefix}01_submission_manuscript_v1.docx`
- Title page: `{manuscript_prefix}02_title_page_v1.docx`
- Cover letter: `{manuscript_prefix}03_cover_letter_v1.docx`

## Figures

- Figure 1: `{figure_prefix}figure1_study_workflow.pdf`
- Figure 2: `{figure_prefix}figure2_trends_breakpoints.pdf`
- Figure 3: `{figure_prefix}figure3_decomposition.pdf`
- Figure 4: `{figure_prefix}figure4_vulnerability.pdf`

## Additional Files

- Additional file 1: `{supplemental_prefix}breakpoint_model_comparison_1990_2021.csv`
- Additional file 2: `{supplemental_prefix}breakpoint_model_comparison_1990_2019.csv`
- Additional file 3: `{supplemental_prefix}four_factor_decomposition_all_periods_by_sex.csv`
- Additional file 4: `{supplemental_prefix}four_factor_decomposition_capped_paf_both.csv`
- Additional file 5: `{supplemental_prefix}four_factor_decomposition_annual_chain_both.csv`
- Additional file 6: `{supplemental_prefix}four_factor_decomposition_common_age_set_both.csv`
- Additional file 7: `{supplemental_prefix}four_factor_decomposition_sensitivity_both.csv`
- Additional file 8: `{supplemental_prefix}vulnerability_focus_long.csv`
- Additional file 9: `{supplemental_prefix}vulnerability_peak_age_long.csv`
- Additional file 10: `{supplemental_prefix}{resolved_repro_bundle_name}`

Do not upload the local helper file below; it is included only to prevent numbering mistakes during manual submission.

- Local helper: `{supplemental_prefix}ADDITIONAL_FILES_UPLOAD_MAP.md`
"""
    (dst / "README.md").write_text(text, encoding="utf-8")


def write_repro_bundle_readme(dst: Path, bundle_name: str, *, current: bool) -> None:
    version_label = bundle_name.rsplit("_", 1)[-1]
    if current:
        summary_line = (
            "This bundle is a non-archival convenience output for the current manuscript-"
            "packaging workflow. It preserves the original project tree for the files needed to "
            "audit the submitted package and regenerate manuscript-facing tables, figures, and "
            "the submission package from the bundled audit inputs."
        )
    else:
        summary_line = (
            "This bundle preserves the original project tree for the files needed to audit the "
            "submitted package and regenerate manuscript-facing tables, figures, and the "
            "submission package from the bundled audit inputs."
        )
    text = f"""# Reproducibility Bundle {version_label}

{summary_line}

It is not a full raw-data or full analysis-pipeline rerun archive. Raw IHME downloads are not redistributed. The bundled files are sufficient for breakpoint rerun and for manuscript-facing package regeneration, including figure/table rebuilding. Complete harmonization, four-factor decomposition, and vulnerability reruns require upstream raw or intermediate inputs that are not redistributed here.

## Software Metadata

- Project name: East Asia respiratory burden manuscript package
- Project home page: `{PUBLIC_COMPANION_URL}`
- Archived version: Zenodo concept DOI `{ZENODO_DOI}` and journal-hosted `{bundle_name}.zip`
- Operating system: Linux (validated on an Ubuntu-compatible server)
- Programming language: Python 3
- Other requirements: standard scientific Python environment with `pandas`, `matplotlib`, and `python-docx`
- License: Creative Commons Attribution 4.0 International (CC BY 4.0) as stated in the bundle-root `LICENSE` file
- Restrictions for non-academic use: none beyond attribution for the manuscript-companion materials; source-data reuse remains subject to IHME terms

## Rebuild Entry Point

- Main package builder: `python3 scripts/70_build_submission_package.py`
- Expected output root: `submission_packages/`
- The builder creates `submission_packages/` automatically if it does not already exist
- In a clean unzip environment, regenerated package labels restart from `v1/current`; these regenerated labels are non-archival convenience outputs, while the archived peer-review snapshot remains the frozen package pair distributed with this release
- Script-by-script run scope and internal table-to-manuscript mapping are documented in `RUN_STATUS.md`

## Scope Note

This freeze contains only the BMC route and is intended for journal submission and external audit of the BMC package.
"""
    (dst / "README.md").write_text(text, encoding="utf-8")


def zip_directory(src_dir: Path, zip_path: Path) -> None:
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for p in sorted(src_dir.rglob("*")):
            if p.is_file():
                zf.write(p, p.relative_to(SUBMISSION_ROOT))


def build_package(spec: dict[str, object]) -> tuple[Path, Path, Path, Path]:
    upload_name = next_version_name(str(spec["upload_prefix"]))
    repro_name = next_version_name(str(spec["repro_prefix"]))
    current_upload_name = str(spec["current_upload_name"])
    current_repro_name = str(spec["current_repro_name"])
    journal_name = str(spec["journal_name"])
    cover_letter_source = Path(spec["cover_letter_source"])
    package_dir = SUBMISSION_ROOT / upload_name
    zip_path = SUBMISSION_ROOT / f"{upload_name}.zip"
    repro_dir = SUBMISSION_ROOT / repro_name
    repro_zip_path = SUBMISSION_ROOT / f"{repro_name}.zip"
    current_package_dir = SUBMISSION_ROOT / current_upload_name
    current_zip_path = SUBMISSION_ROOT / f"{current_upload_name}.zip"
    current_repro_dir = SUBMISSION_ROOT / current_repro_name
    current_repro_zip_path = SUBMISSION_ROOT / f"{current_repro_name}.zip"
    repro_sources = [
        WRITING / "manuscript_full_draft_v10.md",
        WRITING / "abstract_main_text_draft_v3.md",
        WRITING / "title_package_v2.md",
        WRITING / "data_code_availability_statements_v1.md",
        WRITING / "figure_table_captions.md",
        WRITING / "cover_letter_bmc_v1.md",
        PROJECT_ROOT / "data_processed" / "analysis_ready" / "national_summary_panel.csv",
        OUTPUT_MAIN / "breakpoint_series_summary_1990_2021.csv",
        OUTPUT_MAIN / "breakpoint_segments_1990_2021.csv",
        OUTPUT_MAIN / "table2_breakpoint_main.csv",
        OUTPUT_MAIN / "table3_decomposition_main.csv",
        OUTPUT_MAIN / "four_factor_decomposition_main_both.csv",
        OUTPUT_MAIN / "vulnerability_main_snapshot.csv",
        OUTPUT_MAIN / "figure1_study_workflow.png",
        OUTPUT_MAIN / "figure1_study_workflow.pdf",
        OUTPUT_MAIN / "figure2_trends_breakpoints.png",
        OUTPUT_MAIN / "figure2_trends_breakpoints.pdf",
        OUTPUT_MAIN / "figure3_decomposition.png",
        OUTPUT_MAIN / "figure3_decomposition.pdf",
        OUTPUT_MAIN / "figure4_vulnerability.png",
        OUTPUT_MAIN / "figure4_vulnerability.pdf",
        OUTPUT_SUPP / "table3_decomposition_full_with_percent.csv",
        OUTPUT_SUPP / "breakpoint_model_comparison_1990_2021.csv",
        OUTPUT_SUPP / "breakpoint_model_comparison_1990_2019.csv",
        OUTPUT_SUPP / "four_factor_decomposition_all_periods_by_sex.csv",
        OUTPUT_SUPP / "four_factor_decomposition_sensitivity_both.csv",
        OUTPUT_SUPP / "four_factor_decomposition_capped_paf_both.csv",
        OUTPUT_SUPP / "four_factor_decomposition_annual_chain_both.csv",
        OUTPUT_SUPP / "four_factor_decomposition_common_age_set_both.csv",
        OUTPUT_SUPP / "vulnerability_focus_long.csv",
        OUTPUT_SUPP / "vulnerability_peak_age_long.csv",
        LOGS / "four_factor_decomposition_summary.json",
        PROJECT_ROOT / "metadata" / "gbd_download_manifest.csv",
        PROJECT_ROOT / "metadata" / "gbd_download_schema.md",
        PROJECT_ROOT / "LICENSE",
        PROJECT_ROOT / "scripts" / "00_validate_gbd_downloads.py",
        PROJECT_ROOT / "scripts" / "20_harmonize_national_gbd.py",
        PROJECT_ROOT / "scripts" / "21_prepare_four_factor_inputs.py",
        PROJECT_ROOT / "scripts" / "30_breakpoint_trends.py",
        PROJECT_ROOT / "scripts" / "41_decompose_national_burden_four_factor.py",
        PROJECT_ROOT / "scripts" / "50_profile_vulnerability.py",
        PROJECT_ROOT / "scripts" / "60_make_publication_tables.py",
        PROJECT_ROOT / "scripts" / "61_make_publication_figures.py",
        PROJECT_ROOT / "scripts" / "70_build_submission_package.py",
    ]

    build_current_repro_bundle(repro_sources, repro_dir, repro_name, current=False)
    zip_directory(repro_dir, repro_zip_path)

    build_current_repro_bundle(repro_sources, current_repro_dir, current_repro_name, current=True)
    zip_directory(current_repro_dir, current_repro_zip_path)

    build_upload_bundle(package_dir, upload_name, journal_name, cover_letter_source, repro_zip_path)
    zip_directory(package_dir, zip_path)

    build_upload_bundle(current_package_dir, current_upload_name, journal_name, cover_letter_source, current_repro_zip_path)
    zip_directory(current_package_dir, current_zip_path)

    return (
        package_dir,
        zip_path,
        repro_dir,
        repro_zip_path,
        current_package_dir,
        current_zip_path,
        current_repro_dir,
        current_repro_zip_path,
    )


def main() -> int:
    built = []
    for spec in PACKAGE_SPECS:
        built.append(build_package(spec))

    print("OK: submission packages built")
    for (
        package_dir,
        zip_path,
        repro_dir,
        repro_zip_path,
        current_package_dir,
        current_zip_path,
        current_repro_dir,
        current_repro_zip_path,
    ) in built:
        print(" -", package_dir)
        print(" -", zip_path)
        print(" -", repro_dir)
        print(" -", repro_zip_path)
        print(" -", current_package_dir)
        print(" -", current_zip_path)
        print(" -", current_repro_dir)
        print(" -", current_repro_zip_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
